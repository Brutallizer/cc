# -*- coding: utf-8 -*-
"""
Script untuk membuat file jurnal .docx berformat standar IEEE DUA KOLOM
Menggunakan python-docx
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table_row(table, cells_data, bold=False, header=False):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(7)
        run.font.name = 'Times New Roman'
        run.bold = bold
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if header:
            set_cell_shading(cell, "2F5496")
            run.font.color.rgb = RGBColor(255, 255, 255)

def set_columns(section, num_cols, spacing=Cm(0.63)):
    """Set the number of columns for a section."""
    sectPr = section._sectPr
    # Remove existing cols element if any
    for cols in sectPr.findall(qn('w:cols')):
        sectPr.remove(cols)
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), str(num_cols))
    if num_cols > 1:
        cols.set(qn('w:space'), str(int(spacing)))
    sectPr.append(cols)

def add_section_break(doc, break_type='continuous'):
    """Add a section break (continuous) to switch column layouts."""
    new_section = doc.add_section()
    # Set section type via XML directly
    sectPr = new_section._sectPr
    type_elem = sectPr.find(qn('w:type'))
    if type_elem is None:
        type_elem = OxmlElement('w:type')
        sectPr.insert(0, type_elem)
    type_elem.set(qn('w:val'), 'continuous')
    # Copy margins from first section
    first = doc.sections[0]
    new_section.page_width = first.page_width
    new_section.page_height = first.page_height
    new_section.top_margin = first.top_margin
    new_section.bottom_margin = first.bottom_margin
    new_section.left_margin = first.left_margin
    new_section.right_margin = first.right_margin
    return new_section

def create_jurnal():
    doc = Document()
    
    # ===== PAGE SETUP (Section 1: Single Column for Title/Abstract) =====
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(1.91)
    section.right_margin = Cm(1.91)

    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)

    # ===== JUDUL (SINGLE COLUMN) =====
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Rancang Bangun Sistem Verifikasi Keaslian Dokumen Ijazah Berbasis Smart Contract Menggunakan Kriptografi SHA-256 pada Jaringan Polygon')
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    title.space_after = Pt(6)

    # PENULIS
    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author.add_run('Khalid Hanif Albarry')
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    afiliasi = doc.add_paragraph()
    afiliasi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = afiliasi.add_run('Jurusan Ilmu Komputer, Fakultas Komputer, Universitas Insan Cita Indonesia\nJln. Swadaya, Kota Tangerang, 14154, Indonesia\nEmail: khalidhnf1@gmail.com')
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    run.italic = True
    afiliasi.space_after = Pt(10)

    # ABSTRACT (English)
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run('Abstract')
    run.bold = True; run.italic = True; run.font.size = Pt(9)

    abs_en = doc.add_paragraph()
    abs_en.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = abs_en.add_run(
        'Academic document verification in Indonesia remains plagued by slow manual procedures, high legalization costs, and an alarming rise in diploma forgery cases. This research aims to design and develop a blockchain-based diploma verification system called CertiChain, leveraging Smart Contract technology on the Ethereum Virtual Machine (EVM) deployed to the Polygon network. The system employs a hybrid architecture: client-side SHA-256 hashing via the Web Crypto API to generate cryptographic fingerprints of student graduation data, and on-chain storage through a Solidity v0.8.24 Smart Contract connected via Ethers.js v6. A key innovation is the B2B Approval Workflow, where institutions must register their cryptographic identity (wallet address) and receive approval from a designated government authority (Super Admin) before gaining write access to the blockchain ledger. Testing was conducted using Hardhat\'s TDD framework with Mocha and Chai, covering 15 test scenarios. All 15 scenarios achieved a 100% pass rate, confirming the system\'s functional reliability.'
    )
    run.font.size = Pt(9); run.italic = True
    abs_en.space_after = Pt(2)

    kw = doc.add_paragraph()
    r1 = kw.add_run('Keywords: '); r1.bold = True; r1.italic = True; r1.font.size = Pt(8)
    r2 = kw.add_run('Blockchain, Smart Contract, Diploma Verification, SHA-256, Polygon, Solidity, Immutability, Decentralization, B2B Approval')
    r2.italic = True; r2.font.size = Pt(8)
    kw.space_after = Pt(6)

    # ABSTRAK (Indonesia)
    h2 = doc.add_paragraph()
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h2.add_run('Abstrak')
    run.bold = True; run.italic = True; run.font.size = Pt(9)

    abs_id = doc.add_paragraph()
    abs_id.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = abs_id.add_run(
        'Proses verifikasi keaslian dokumen akademik di Indonesia masih menghadapi kendala berupa prosedur manual yang lambat, biaya legalisir yang tinggi, serta meningkatnya kasus pemalsuan ijazah. Penelitian ini merancang sistem verifikasi bernama CertiChain berbasis Blockchain dan Smart Contract. Sistem dikembangkan menggunakan Solidity v0.8.24 pada jaringan Polygon, dengan antarmuka HTML5, Vanilla JavaScript, TailwindCSS, dan Ethers.js v6. Inovasi utama adalah arsitektur B2B Approval Workflow di mana institusi harus mendapat persetujuan Kementerian sebelum menerbitkan data. Hasil pengujian terhadap 15 skenario uji menunjukkan success rate 100%.'
    )
    run.font.size = Pt(9); run.italic = True
    abs_id.space_after = Pt(2)

    kw2 = doc.add_paragraph()
    r1 = kw2.add_run('Kata Kunci: '); r1.bold = True; r1.italic = True; r1.font.size = Pt(8)
    r2 = kw2.add_run('Blockchain, Smart Contract, Verifikasi Ijazah, SHA-256, Polygon, Solidity, Immutability, Desentralisasi')
    r2.italic = True; r2.font.size = Pt(8)
    kw2.space_after = Pt(8)

    # ===== SECTION BREAK: Switch to TWO COLUMNS =====
    sec2 = add_section_break(doc, 'continuous')
    set_columns(sec2, 2, spacing=Cm(0.63))

    # ===== HELPERS =====
    def add_heading_styled(text, level=1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10) if level == 1 else Pt(9)
        run.font.name = 'Times New Roman'
        p.space_before = Pt(8)
        p.space_after = Pt(4)
        return p

    def add_body(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        p.paragraph_format.first_line_indent = Cm(0.5)
        p.space_after = Pt(3)
        p.space_before = Pt(0)
        return p

    # ===== I. PENDAHULUAN =====
    add_heading_styled('I. PENDAHULUAN', 1)

    add_heading_styled('1.1 Latar Belakang Umum', 2)
    add_body('Verifikasi keaslian dokumen akademik merupakan proses krusial yang menghubungkan dunia pendidikan tinggi dengan dunia industri, khususnya dalam proses rekrutmen tenaga kerja profesional. Di Indonesia, proses ini masih sangat bergantung pada mekanisme konvensional seperti legalisir stempel basah, pengiriman surat verifikasi antar-universitas, hingga pengecekan manual melalui Pangkalan Data Pendidikan Tinggi (PDDIKTI) yang kerap mengalami keterlambatan respons atau downtime server. Perkembangan teknologi Blockchain menawarkan paradigma baru berupa penyimpanan data yang bersifat desentralisasi, transparan, dan tidak dapat diubah (immutable), menjadikannya kandidat solusi ideal untuk permasalahan integritas dokumen. Penelitian ini mengusulkan pembangunan sistem verifikasi keaslian ijazah bernama CertiChain yang memanfaatkan Smart Contract pada jaringan Polygon untuk menyimpan sidik jari kriptografi dari data kelulusan mahasiswa.')

    add_heading_styled('1.2 Latar Belakang Spesifik', 2)
    add_body('Isu pemalsuan ijazah di era digital telah menjadi permasalahan serius yang mengancam kredibilitas institusi pendidikan tinggi dan merugikan pelamar kerja yang jujur. Sistem tata kelola verifikasi yang ada saat ini, baik yang bersifat manual maupun terdigitalisasi parsial menggunakan database terpusat, memiliki kerentanan fundamental berupa Single Point of Failure. Kelemahan paling fatal dari sistem tersentralisasi bukanlah sekadar ancaman peretas dari luar, melainkan kerentanan terhadap Insider Threat (Ancaman Orang Dalam). Administrator database memiliki wewenang penuh yang memungkinkan mereka memanipulasi, menyisipkan, atau menghapus data kelulusan tanpa meninggalkan jejak audit yang transparan. Kondisi ini menunjukkan urgensi kebutuhan akan sistem yang bersifat trustless dan tamper-proof. Arsitektur Blockchain diusulkan karena menawarkan sifat Trustless Integrity \u2014 tidak ada satu entitas pun yang dapat mengubah data yang telah dikonfirmasi oleh konsensus jaringan.')

    add_heading_styled('1.3 Kesenjangan Pengetahuan', 2)
    add_body('Sebagian besar solusi administrasi akademik digital hingga kini masih terbatas pada digitalisasi parsial, seperti penyimpanan sertifikat di server cloud atau penerapan kode QR statis. Beberapa penelitian terdahulu mengusulkan penggunaan Blockchain untuk sertifikasi digital, namun mayoritas berfokus pada penyimpanan file utuh yang menimbulkan gas fee sangat tinggi. Selain itu, belum ada yang mengimplementasikan mekanisme otorisasi bertingkat (multi-tier approval) yang mereplikasi birokrasi penerbitan ijazah di dunia nyata. Penelitian ini mengisi kesenjangan tersebut dengan Zero-Knowledge Hashing dan B2B Approval Workflow.')

    add_heading_styled('1.4 Kontribusi Penelitian', 2)
    add_body('Penelitian ini mengusulkan arsitektur dApp CertiChain yang mentransformasi verifikasi dokumen dari paradigma tersentralisasi menjadi terdesentralisasi. Proses hashing SHA-256 dilakukan di sisi klien via Web Crypto API, sementara penyematan hash ke blockchain dilakukan melalui Smart Contract Solidity. Kontribusi tambahan adalah implementasi B2B Approval Workflow di mana institusi harus melalui pendaftaran mandiri dan mendapat persetujuan Kementerian sebelum dapat menerbitkan ijazah digital.')

    # ===== II. PENELITIAN TERKAIT =====
    add_heading_styled('II. PENELITIAN YANG TERKAIT', 1)
    add_body('Turkanovi\u0107 et al. (2018) mengembangkan EduCTX untuk transfer kredit akademik berbasis metadata on-chain [9]. Arenas dan Fernandez (2018) mengusulkan CredenceLedger berbasis permissioned blockchain [10]. Wibowo dkk. (2020) mengujicobakan blockchain privat untuk metadata ijazah [11]. Zheng et al. (2017) memberikan tinjauan arsitektur blockchain dan mekanisme konsensus [12]. Grech dan Camilleri (2017) mengidentifikasi potensi blockchain dalam pendidikan namun menekankan tantangan skalabilitas [8]. CertiChain menyempurnakan penelitian terdahulu dengan Zero-Knowledge Hashing pada jaringan publik Polygon serta lapisan B2B Approval yang belum ada sebelumnya.')

    # ===== III. METODE =====
    add_heading_styled('III. METODE PENELITIAN', 1)

    add_heading_styled('3.1 Pendekatan Pengembangan', 2)
    add_body('Pengembangan dApp CertiChain menggunakan metode Iterative Incremental Development. Fase V1 berfokus pada fungsi inti penyimpanan dan verifikasi hash. Fase V2 menambahkan pendaftaran institusi mandiri dan approval workflow oleh Kementerian.')

    add_heading_styled('3.2 Arsitektur Sistem', 2)
    add_body('Arsitektur sistem terdiri dari empat lapisan: Presentation (HTML5, TailwindCSS), Application (JavaScript ES6 untuk hashing dan validasi), Integration (Ethers.js v6 sebagai Web3 Bridge), dan Blockchain (Solidity v0.8.24 untuk penyimpanan hash dan kontrol akses).')

    add_heading_styled('3.3 Alur Kerja B2B Approval', 2)
    add_body('Prosedur operasional CertiChain mereplikasi hierarki birokrasi ke dalam Smart Contract dengan tiga peran: (1) Kementerian (Super Admin) memiliki hak approveInstitution() dan rejectInstitution(). (2) Admin Kampus (status Approved) dapat mengakses storeHash() dan storeMultipleHashes(). (3) Publik/HRD dapat memverifikasi gratis melalui verifyHash().')
    add_body('Alur prosedural: Fase Pengajuan \u2014 institusi memanggil applyForRegistration(), status menjadi Pending. Fase Otorisasi \u2014 Kementerian mengeksekusi approveInstitution(). Fase Penerbitan \u2014 kampus Approved menyimpan hash SHA-256 ke blockchain. Fase Verifikasi \u2014 HRD mencocokkan hash dari dokumen asli dengan data on-chain.')

    add_heading_styled('3.4 Desain Smart Contract', 2)
    add_body('Smart Contract V2 menggunakan struct Institution (name, status enum), mapping institutions, array applicantList, serta modifier onlyApprovedInstitution yang menjamin hanya institusi sah yang dapat menulis ke ledger. Fungsi verifyHash() bersifat view (gratis) dan mengembalikan status validitas, nama institusi, dan alamat publisher.')

    # ===== IV. HASIL =====
    add_heading_styled('IV. HASIL DAN PEMBAHASAN', 1)

    add_heading_styled('4.1 Implementasi Antarmuka', 2)
    add_body('Sistem menghasilkan empat antarmuka: (1) Portal Pendaftaran Institusi (register.html), (2) Dashboard Kementerian untuk Approve/Reject, (3) Dashboard Admin Kampus dengan form manual, Bulk Import CSV, dan riwayat transaksi, (4) Portal Verifikasi HRD (verify.html) yang menampilkan profil kampus penerbit saat hash valid.')

    add_heading_styled('4.2 Pengujian Smart Contract', 2)
    add_body('Pengujian dieksekusi menggunakan Mocha dan Chai melalui Hardhat Network dengan pendekatan TDD. Seluruh 15 skenario mencatat status PASS: deployment, registration (mandiri + duplikat), approval, rejection, access control (non-admin approve, pending store, non-approved store), hash storage (tunggal + duplikat), bulk store (normal + duplikat), dan verification (valid, palsu, publik).')
    add_body('Success rate mencapai 100%, membuktikan Smart Contract berfungsi tanpa cacat kerentanan pada seluruh skenario fungsional dan keamanan.')

    add_heading_styled('4.3 Deployment Jaringan Publik', 2)
    add_body('Smart Contract V2 berhasil di-deploy pada Polygon Amoy Testnet dengan alamat kontrak 0xF20276816FDEb9f76Bd385086CEB8e44826B689b. Pengujian End-to-End pada jaringan publik mengkonfirmasi konsistensi fungsi dengan pengujian lokal, memvalidasi kesiapan migrasi ke Polygon Mainnet.')

    # ===== V. KESIMPULAN =====
    add_heading_styled('V. KESIMPULAN', 1)

    conclusions = [
        'Arsitektur CertiChain berhasil membuktikan keandalan sistem verifikasi mandiri berbasis EVM dengan skema gratis bagi Verifikator/HRD.',
        'Kriptografi SHA-256 mencapai target Privacy by Design \u2014 data pribadi mahasiswa tidak tersimpan di blockchain.',
        'Strategi Bulk Import menekan biaya gas fee secara signifikan untuk penerbitan massal.',
        'Logika Smart Contract tervalidasi 15 skenario TDD dengan success rate 100%.',
        'B2B Approval Workflow berhasil menciptakan lapisan otorisasi terdesentralisasi yang memblokir institusi tidak berizin.',
    ]
    for i, c in enumerate(conclusions):
        p = doc.add_paragraph()
        run = p.add_run(f'{i+1}. {c}')
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        p.space_after = Pt(2)

    # ===== UCAPAN TERIMA KASIH =====
    add_heading_styled('UCAPAN TERIMA KASIH', 1)
    add_body('Penulis mengucapkan terima kasih kepada dosen pembimbing dan civitas akademika Universitas Insan Cita Indonesia (UICI) atas dukungan dan bimbingan selama penelitian ini berlangsung.')

    # ===== DAFTAR PUSTAKA =====
    add_heading_styled('DAFTAR PUSTAKA', 1)
    refs = [
        '[1] V. Buterin, "A Next-Generation Smart Contract and Decentralized Application Platform," Ethereum White Paper, 2014.',
        '[2] Ethers.js Documentation, "Ethers.js v6," 2024. [Online]. Available: https://docs.ethers.org/v6/',
        '[3] Hardhat Documentation, "Ethereum Development Environment," 2024. [Online]. Available: https://hardhat.org/docs',
        '[4] S. Nakamoto, "Bitcoin: A Peer-to-Peer Electronic Cash System," 2008.',
        '[5] NIST, "FIPS PUB 180-2: Secure Hash Standard (SHS)," U.S. DOC, 2001.',
        '[6] OpenZeppelin, "Smart Contract Security Best Practices," 2024.',
        '[7] Solidity Documentation, "Solidity v0.8.x," 2024. [Online]. Available: https://docs.soliditylang.org/',
        '[8] A. Grech and A. F. Camilleri, "Blockchain in Education," JRC, European Commission, 2017.',
        u'[9] M. Turkanovi\u0107 et al., "EduCTX: A Blockchain-Based Higher Education Credit Platform," IEEE Access, vol. 6, pp. 5112-5127, 2018.',
        '[10] R. Arenas and P. Fernandez, "CredenceLedger," IEEE ICETI, 2018.',
        '[11] S. A. Wibowo et al., "Prototyping e-Degree Management System Using Blockchain," JSINBIS, vol. 10, no. 2, pp. 231-238, 2020.',
        '[12] Z. Zheng et al., "An Overview of Blockchain Technology," IEEE ICBD, pp. 557-564, 2017.',
        '[13] G. Wood, "Ethereum: A Secure Decentralised Generalised Transaction Ledger," Yellow Paper, 2014.',
    ]
    for ref in refs:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        run.font.size = Pt(7)
        run.font.name = 'Times New Roman'
        p.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(0.35)
        p.paragraph_format.first_line_indent = Cm(-0.35)

    # ===== SAVE =====
    output_path = r"C:\Users\aldip\OneDrive\Desktop\Jurnal_CertiChain_DuaKolom.docx"
    doc.save(output_path)
    print(f"SUKSES! File jurnal DUA KOLOM tersimpan di: {output_path}")

if __name__ == '__main__':
    create_jurnal()
